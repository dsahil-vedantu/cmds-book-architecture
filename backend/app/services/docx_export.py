"""Native python-docx exporter for theory + questions + regen.

Replaces the previous pandoc-based docx pipeline so we can control:
  - exact heading styles, font sizes, paragraph spacing
  - no-duplicate-headings invariant (a `last_heading` tracker silently
    drops adjacent identical headings — covers case/punct differences)
  - per-question label format (Question 1: / Options: / Solution: with
    bold colon-suffixed labels, hanging indent so wrapped lines align)
  - empty fields are SKIPPED entirely (no "Answer:" stub when there's
    no answer key)
  - explicit gaps between sections / questions for predictable layout

Pandoc is no longer required for docx export. (Markdown export still
uses the existing builders in the API routers.)

Public entry points:
  - build_questions_docx(title, sections, section_only=None) -> bytes
  - build_regen_docx(regen_label, custom_instructions, sections, ...) -> bytes
  - build_theory_docx(book_title, sections) -> bytes
"""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.services.latex_omml import latex_to_omml_element
from app.services.latex_normalize import normalize_latex
from app.core.config import settings


# ---------------------------------------------------------------------------
# Shared building blocks
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x1A, 0x36, 0x6E)
KEYPOINT_ORANGE = RGBColor(0xC8, 0x7C, 0x00)
MUTED = RGBColor(0x66, 0x66, 0x66)

# Inline `$...$` and display `$$...$$` LaTeX chunks. python-docx can't
# render OMML equations directly without significant effort, so for v1
# we wrap math chunks in italic runs — preserves placement and reads
# correctly. (Future: write OMML XML directly for native equations.)
MATH_RE = re.compile(r"\$\$?(.+?)\$\$?", re.DOTALL)

# MCQ option marker inside raw_text (case-insensitive, A-D or 1-4 inside
# parens). Used to (a) strip options off the stem so the Question line
# doesn't repeat them, and (b) lay them out cleanly under Options:.
OPTION_RE = re.compile(r"\(([A-D1-4])\)")
# Strict-MCQ pattern: only UPPERCASE A-D or digits 1-4. Lowercase (a)(b)(c)(d),
# roman numerals (i)(ii)(iii)(iv), and any other format are deliberately NOT
# matched — for those formats the raw_text stays intact (stem keeps everything,
# no separate Options block). This avoids the "we partially recognized the
# format and stripped content we shouldn't have" failure mode: when in doubt
# about the option format, render the whole thing as the question body.

# Figure placeholders embedded by the extractor: {{fig: <label> — <caption>}}
FIG_RE = re.compile(r"\{\{\s*fig\s*:\s*([^}]+?)\s*\}\}", re.IGNORECASE)

# Markdown-style tables embedded in raw_text — Gemini emits these for
# value tables (x/y), match-the-column, etc. We detect 2+ consecutive
# lines starting with `|`, where line 2 is a `|---|---|` separator,
# and emit a proper Word table.
_PIPE_LINE = re.compile(r"^\s*\|.*\|\s*$")
_PIPE_DIV = re.compile(r"^\s*\|[\s:|\-]+\|\s*$")

# Common LaTeX commands that leak out of `$...$` chunks because Gemini
# sometimes uses them in prose. Map to Unicode so the docx reads
# naturally instead of showing the raw backslash sequence.
_LATEX_TO_UNICODE: list[tuple[str, str]] = [
    # Order matters — longer prefixes first to avoid `\rightarrow`
    # being half-substituted by `\right`.
    (r"\Leftrightarrow", "⇔"),
    (r"\Rightarrow", "⇒"),
    (r"\Leftarrow", "⇐"),
    (r"\rightarrow", "→"),
    (r"\leftarrow", "←"),
    (r"\therefore", "∴"),
    (r"\because", "∵"),
    (r"\approx", "≈"),
    (r"\equiv", "≡"),
    (r"\infty", "∞"),
    (r"\times", "×"),
    (r"\cdot", "·"),
    (r"\div", "÷"),
    (r"\pm", "±"),
    (r"\mp", "∓"),
    (r"\leq", "≤"),
    (r"\geq", "≥"),
    (r"\neq", "≠"),
    (r"\ne", "≠"),
    (r"\to", "→"),
    # Set theory
    (r"\cup", "∪"),
    (r"\cap", "∩"),
    (r"\subseteq", "⊆"),
    (r"\supseteq", "⊇"),
    (r"\subset", "⊂"),
    (r"\supset", "⊃"),
    (r"\notin", "∉"),
    (r"\in", "∈"),
    (r"\emptyset", "∅"),
    (r"\varnothing", "∅"),
    # Logic / quantifiers
    (r"\forall", "∀"),
    (r"\exists", "∃"),
    (r"\lnot", "¬"),
    (r"\neg", "¬"),
    (r"\land", "∧"),
    (r"\lor", "∨"),
    # More misc
    (r"\partial", "∂"),
    (r"\nabla", "∇"),
    (r"\sum", "∑"),
    (r"\prod", "∏"),
    (r"\int", "∫"),
    (r"\oint", "∮"),
    (r"\ldots", "…"),
    (r"\cdots", "⋯"),
    # Greek letters most commonly seen in math prose
    (r"\alpha", "α"), (r"\beta", "β"), (r"\gamma", "γ"),
    (r"\delta", "δ"), (r"\epsilon", "ε"), (r"\theta", "θ"),
    (r"\lambda", "λ"), (r"\mu", "μ"), (r"\pi", "π"),
    (r"\sigma", "σ"), (r"\phi", "φ"), (r"\omega", "ω"),
    (r"\Delta", "Δ"), (r"\Theta", "Θ"), (r"\Lambda", "Λ"),
    (r"\Sigma", "Σ"), (r"\Phi", "Φ"), (r"\Omega", "Ω"),
    # Common spacing / punctuation
    (r"\,", " "), (r"\;", " "), (r"\:", " "), (r"\!", ""),
]

# Single-char super/subscripts: x^2, x^{n}, x_1, x_{i}. Only handles
# single digits/letters that have Unicode super/sub forms. Multi-char
# expressions like ^{n+1} are left as-is.
_SUP_MAP = str.maketrans({
    "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
    "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
    "+": "⁺", "-": "⁻", "=": "⁼", "(": "⁽", ")": "⁾", "n": "ⁿ", "i": "ⁱ",
})
_SUB_MAP = str.maketrans({
    "0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄",
    "5": "₅", "6": "₆", "7": "₇", "8": "₈", "9": "₉",
    "+": "₊", "-": "₋", "=": "₌", "(": "₍", ")": "₎",
    "n": "ₙ", "i": "ᵢ", "a": "ₐ", "e": "ₑ", "o": "ₒ", "x": "ₓ",
})


def _normalise_math_prose(text: str) -> str:
    """Substitute common LaTeX commands → Unicode and basic
    super/subscripts so the rendered docx reads naturally. Applied to
    text after stripping `$...$` math wrappers (math chunks go through
    this too — the rendered italic stays, just with proper symbols)."""
    # Structural LaTeX commands that don't have a single Unicode codepoint
    # but can be rewritten plainly. Done BEFORE the simple substitutions so
    # any greek letters inside survive.
    #   \frac{a}{b}        → (a)/(b)         (parens added when needed)
    #   \dfrac, \tfrac     → same as \frac
    #   \sqrt[n]{x}        → ⁿ√(x)
    #   \sqrt{x}           → √(x)
    #   \text{x}, \mathrm{x}, \mathbf{x}, \operatorname{x} → x
    #   \left, \right      → dropped (display-size hints only)
    def _drop_text(m: re.Match[str]) -> str:
        return m.group(1)

    text = re.sub(r"\\(?:text|mathrm|mathbf|mathit|operatorname)\{([^{}]*)\}", _drop_text, text)
    text = re.sub(r"\\left\b", "", text)
    text = re.sub(r"\\right\b", "", text)

    def _frac(m: re.Match[str]) -> str:
        num = m.group(1).strip()
        den = m.group(2).strip()
        wrap = lambda s: s if len(s) == 1 and s.isalnum() else f"({s})"
        return f"{wrap(num)}/{wrap(den)}"

    def _sqrt(m: re.Match[str]) -> str:
        return f"√({m.group(1)})"

    def _nthroot(m: re.Match[str]) -> str:
        n = m.group(1)
        body = m.group(2)
        sup = {"2": "²", "3": "³", "4": "⁴", "5": "⁵"}.get(n.strip(), n.strip())
        return f"{sup}√({body})"

    # super/subscript helpers (defined here so they can run INSIDE the
    # fixpoint loop below — running them first unblocks sqrt/frac when
    # the args contain `x^{2}` style braces).
    def _sup_inner(m: re.Match[str]) -> str:
        body = m.group(1) or m.group(2) or ""
        if body and all(ord(ch) in _SUP_MAP for ch in body):
            return body.translate(_SUP_MAP)
        return m.group(0)

    def _sub_inner(m: re.Match[str]) -> str:
        body = m.group(1) or m.group(2) or ""
        if body and all(ord(ch) in _SUB_MAP for ch in body):
            return body.translate(_SUB_MAP)
        return m.group(0)

    # Single fixpoint loop covering super/subscripts + sqrt + frac.
    # ORDER MATTERS: super/subscripts must run FIRST so `\sqrt{b^{2}-4ac}`
    # becomes `\sqrt{b²-4ac}` (no inner braces) and the sqrt regex
    # (`[^{}]*`) can then match it. Same logic for frac with `\sqrt`
    # inside. Iterating to fixpoint handles arbitrary nesting.
    for _ in range(8):
        prev = text
        # 1. single-char super/subscripts (e.g. b^{2} → b²)
        text = re.sub(
            r"\^\{([^{}]{1,4})\}|\^([0-9a-zA-Z+\-=()])", _sup_inner, text
        )
        text = re.sub(
            r"_\{([^{}]{1,4})\}|_([0-9a-zA-Z+\-=()])", _sub_inner, text
        )
        # 2. roots
        text = re.sub(r"\\sqrt\[([^\]]+)\]\{([^{}]*)\}", _nthroot, text)
        text = re.sub(r"\\sqrt\{([^{}]*)\}", _sqrt, text)
        # 3. fractions
        text = re.sub(r"\\(?:d|t)?frac\{([^{}]*)\}\{([^{}]*)\}", _frac, text)
        if text == prev:
            break

    for tex, uni in _LATEX_TO_UNICODE:
        text = text.replace(tex, uni)

    # super/subscript handling is now done INSIDE the fixpoint loop above
    # — needs to run before sqrt/frac so inner braces from `x^{2}` don't
    # block the outer regexes. Keeping a second pass here is redundant.

    # Final XML safety pass — python-docx rejects control characters that
    # XML 1.0 disallows. JSON-escape collisions in the source data leave
    # behind \v / \f / \b / \x01-\x08 / \x0E-\x1F etc. These are usually
    # the tail of stripped LaTeX commands (e.g. `\vec` → `\x0B + ec`).
    # Strip them so the docx renders cleanly.
    return _sanitize_xml(text)


# XML 1.0 valid chars: \t \n \r and >= \x20 (except surrogates / FFFE / FFFF)
_XML_INVALID_RE = re.compile(
    r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]"
)

def _sanitize_xml(text: str) -> str:
    """Strip characters that python-docx / XML 1.0 can't accept. Replaces
    them with a single space so the surrounding text stays readable."""
    if not text:
        return text
    return _XML_INVALID_RE.sub(" ", text)


def _extract_tables_from_text(text: str):
    """Yield ('text', str) or ('table', (headers, rows)) tuples.

    Detects a Markdown table as:
        | header1 | header2 | ...|
        |---------|---------|----|
        | r1c1    | r1c2    | ...|
        | r2c1    | r2c2    | ...|
    Anything between tables (or before/after) is yielded as 'text'.
    """
    lines = text.split("\n")
    i = 0
    buf: list[str] = []
    while i < len(lines):
        line = lines[i]
        # Detect start: this line is a pipe row AND next line is a divider
        if (
            _PIPE_LINE.match(line)
            and i + 1 < len(lines)
            and _PIPE_DIV.match(lines[i + 1])
        ):
            # Flush prose buffer first
            if buf:
                yield ("text", "\n".join(buf).strip("\n"))
                buf = []
            # Parse headers (split on |, trim, drop empty leading/trailing)
            def _cells(row: str) -> list[str]:
                return [c.strip() for c in row.strip().strip("|").split("|")]
            headers = _cells(line)
            rows: list[list[str]] = []
            j = i + 2
            while j < len(lines) and _PIPE_LINE.match(lines[j]):
                rows.append(_cells(lines[j]))
                j += 1
            yield ("table", (headers, rows))
            i = j
            continue
        buf.append(line)
        i += 1
    if buf:
        yield ("text", "\n".join(buf).strip("\n"))


def _norm(s: str) -> str:
    """Loose comparison key for heading de-dup — lowercased, alphanumerics
    only. Catches case and punctuation differences."""
    return re.sub(r"[^a-z0-9]+", "", (s or "").lower())


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


# Inline markdown: **bold**, __bold__, \textbf{}, *italic*, \textit/\emph{},
# `code`. Math ($...$) is handled separately by _render_inline before this.
_MD_INLINE = re.compile(
    r"\*\*(?P<b>.+?)\*\*"
    r"|__(?P<b2>.+?)__"
    r"|\\textbf\{(?P<b3>[^}]*)\}"
    r"|(?<![\w*])\*(?P<i>[^*\s][^*]*?)\*(?![\w*])"
    r"|\\(?:textit|emph)\{(?P<i2>[^}]*)\}"
    r"|`(?P<code>[^`]+)`"
)


def _emit_text_runs(p, text: str, *, bold_all: bool = False) -> None:
    """Emit runs for a plain-text segment, rendering inline markdown
    (**bold**, *italic*, `code`, \\textbf/\\textit) as real Word formatting
    instead of leaking the literal markers (** , `, \\textbf{}) into the doc.
    ``bold_all`` forces every run bold (used for leaked heading lines)."""
    pos = 0
    for m in _MD_INLINE.finditer(text):
        if m.start() > pos:
            r = p.add_run(_normalise_math_prose(text[pos:m.start()]))
            r.font.size = Pt(10); r.bold = r.bold or bold_all
        bold = m.group("b") or m.group("b2") or m.group("b3")
        ital = m.group("i") or m.group("i2")
        code = m.group("code")
        if bold is not None:
            r = p.add_run(_normalise_math_prose(bold)); r.bold = True
        elif ital is not None:
            r = p.add_run(_normalise_math_prose(ital)); r.italic = True
            if bold_all:
                r.bold = True
        else:
            r = p.add_run(code); r.font.name = "Consolas"
            if bold_all:
                r.bold = True
        r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(_normalise_math_prose(text[pos:]))
        r.font.size = Pt(10); r.bold = r.bold or bold_all


def _render_inline(p, text: str) -> None:
    """Add inline runs to paragraph p. Math chunks render italic;
    figure placeholders render as muted bracketed callouts."""
    # Defensive XML safety — strip any control chars before passing
    # anything to python-docx (which raises ValueError on them). Also
    # caught by _normalise_math_prose's tail, but inline paths can
    # bypass that (e.g. table cells) so we do it here too.
    text = _sanitize_xml(text or "")
    # Wrap bare / partially-delimited LaTeX math (incl. \ce) in $...$ so the
    # OMML pass below renders it — covers un-delimited math anywhere,
    # including table cells. Idempotent; preserves existing $...$ and
    # figure placeholders. Mirrors the frontend normalizeLatex.
    text = normalize_latex(text)
    # A markdown heading marker (#, ##, ###) leaking into an inline context
    # (e.g. "## Solution" inside a solution body that paragraph() didn't
    # split) — strip the markers and bold the line so no literal #'s reach
    # the doc. Block-level headings are handled earlier in paragraph().
    _force_bold = False
    _hm = re.match(r"^\s*#{1,6}\s+(.*\S)\s*$", text)
    if _hm:
        text = _hm.group(1)
        _force_bold = True
    elif re.search(r"(?m)^[ \t]*#{1,6}[ \t]+", text):
        # Multi-line body with a leaked heading marker (e.g. "## Solution\n…")
        # that paragraph()/labeled() didn't split — strip the leading #'s per
        # line so no literal hashes reach the doc (context supplies the heading).
        text = re.sub(r"(?m)^[ \t]*#{1,6}[ \t]+", "", text)
    # First substitute figure placeholders → bracketed callouts (still
    # processed inline so they stay in flow with surrounding prose).
    parts: list[tuple[str, str]] = []  # (kind, content) kind in {text, math, fig}
    cursor = 0
    # Combined scan: math has priority over figure (math placeholders are
    # syntactically tighter; figures live in plain prose).
    spans: list[tuple[int, int, str, str]] = []  # (start, end, kind, payload)
    for m in MATH_RE.finditer(text):
        spans.append((m.start(), m.end(), "math", m.group(1)))
    for m in FIG_RE.finditer(text):
        # Skip if inside a math span
        if any(s <= m.start() < e for s, e, _, _ in spans):
            continue
        spans.append((m.start(), m.end(), "fig", m.group(1).strip()))
    spans.sort()
    for s, e, kind, payload in spans:
        if s > cursor:
            parts.append(("text", text[cursor:s]))
        parts.append((kind, payload))
        cursor = e
    if cursor < len(text):
        parts.append(("text", text[cursor:]))

    # Bold state is tracked ACROSS parts: a markdown bold span that wraps a
    # `$math$` chunk (e.g. "**Case (1): $S_1$ closed:**") would otherwise be
    # severed by the math-split above, leaking literal `**`. We split each text
    # part on `**` and toggle bold at every marker, so the span stays bold
    # across the math boundary and no stray stars reach the doc. Within each
    # piece, _emit_text_runs still handles *italic*, `code`, and \textbf{}.
    bold_state = _force_bold
    for kind, payload in parts:
        if kind == "text":
            segs = payload.split("**")
            for i, seg in enumerate(segs):
                if seg:
                    _emit_text_runs(p, seg, bold_all=bold_state)
                if i < len(segs) - 1:
                    bold_state = not bold_state  # toggle at each **
        elif kind == "math":
            # Native Word equation (OMML) — fractions, integrals, roots,
            # matrices, \ce reactions render properly, matching the preview.
            omath = latex_to_omml_element(payload)
            if omath is not None:
                p._p.append(omath)
            else:
                # Fallback: legacy unicode-approximation for the long tail
                # of LaTeX the converter can't handle (never crash the doc).
                r = p.add_run(_normalise_math_prose(payload))
                r.italic = True
                r.bold = bold_state
                r.font.size = Pt(10)
        elif kind == "fig":
            # Strip the placeholder silently. The actual figure is rendered
            # separately below the question text via the embedded_figures
            # pipeline, so an inline "[Figure: ...]" muted callout here is
            # redundant. Keeping the callout caused visible duplication
            # ("[Figure: A, B, C, D - (unlabelled diagram)]" + image of
            # the same figure rendered right below).
            pass


def _strip_options_from_stem(raw_text: str) -> str:
    """Return the question text with the trailing option block removed."""
    m = OPTION_RE.search(raw_text)
    if not m:
        return raw_text.strip()
    return raw_text[: m.start()].strip()


def _parse_options(raw_text: str) -> list[tuple[str, str]]:
    """Split out lettered options from raw_text. Returns [(letter, body)]."""
    parts = OPTION_RE.split(raw_text)
    if len(parts) < 5:
        return []
    opts: list[tuple[str, str]] = []
    # parts is [stem, 'A', text_A, 'B', text_B, …]; iterate over pairs
    for i in range(1, len(parts) - 1, 2):
        letter = parts[i].upper()
        body = parts[i + 1].strip()
        # Strip an orphan opening paren left from the next letter
        body = body.rstrip("(").rstrip().rstrip(",").strip()
        opts.append((letter, body))
    return opts


def _split_answer_from_solution(sol: str) -> tuple[str, str]:
    """If the printed solution begins with 'Ans.' / 'Answer:' / '(B)',
    pull that first line off as the Answer; rest is the Solution."""
    sol = (sol or "").strip()
    if not sol:
        return "", ""
    first_line = sol.split("\n", 1)[0].strip()
    if re.match(r"^(Ans\.?\:?|Answer\:?\.?|\([A-Da-d]\))", first_line, re.I):
        rest = sol[len(first_line):].lstrip("\n .:")
        return first_line, rest
    return "", sol


# ---------------------------------------------------------------------------
# Heading + spacing helpers (with de-dup guard)
# ---------------------------------------------------------------------------

def _parse_latex_tabular(text: str):
    """Parse a LaTeX ``\\begin{tabular}{...} ... \\end{tabular}`` block into
    (headers, rows). Cells split on ``&``, rows on ``\\\\``; rules (\\hline,
    booktabs) are dropped. First row becomes the header. Returns
    (None, None) if no tabular is present (caller leaves the block as-is)."""
    if not text or "\\begin{tabular}" not in text:
        return None, None
    m = re.search(
        r"\\begin\{tabular\}(?:\{[^}]*\})?(.*?)\\end\{tabular\}", text, re.DOTALL
    )
    if not m:
        return None, None
    body = re.sub(
        r"\\hline|\\toprule|\\midrule|\\bottomrule|\\cline\{[^}]*\}", "", m.group(1)
    )
    raw_rows = [r.strip() for r in re.split(r"\\\\", body) if r.strip()]
    if not raw_rows:
        return None, None
    cells = [[c.strip() for c in r.split("&")] for r in raw_rows]
    return cells[0], cells[1:]


class _DocBuilder:
    """Holds a Document plus a last-heading tracker so de-dup works
    across the whole document."""

    def __init__(self) -> None:
        self.doc = Document()
        _set_default_font(self.doc)
        self._last_heading = ""

    def title(self, text: str) -> None:
        text = _sanitize_xml(text or "")
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(14)
        self.doc.add_paragraph()  # one-line breather

    def group_header(self, text: str) -> None:
        """Big divider — WORKED EXAMPLES, JEE SPECIAL WING, etc."""
        text = _sanitize_xml(text or "")
        if _norm(text) == _norm(self._last_heading):
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = NAVY
        self._last_heading = text

    def section_heading(self, text: str) -> None:
        """Section-level heading — EXAMPLE 4.7, Introduction, etc."""
        text = _sanitize_xml(text or "")
        if _norm(text) == _norm(self._last_heading):
            return
        h = self.doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = NAVY
        self._last_heading = text

    def sub_heading(self, text: str) -> None:
        text = _sanitize_xml(text or "")
        if _norm(text) == _norm(self._last_heading):
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        self._last_heading = text

    def paragraph(self, text: str, *, space_after_pt: int = 3,
                  left_indent_cm: float = 0.0) -> None:
        text = (text or "").strip()
        # Markdown content may carry heading lines (#, ##, ###) and blank-line
        # paragraph breaks. Split on newlines: a `#{1,6} ` line becomes a
        # sub-heading; everything else is a paragraph with inline markdown.
        if "\n" in text or re.match(r"^#{1,6}\s", text):
            for raw in re.split(r"\n+", text):
                line = raw.strip()
                if not line:
                    continue
                hm = re.match(r"^#{1,6}\s+(.+)$", line)
                if hm:
                    self.sub_heading(hm.group(1).strip())
                else:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(space_after_pt)
                    if left_indent_cm:
                        p.paragraph_format.left_indent = Cm(left_indent_cm)
                    _render_inline(p, line)
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after_pt)
        if left_indent_cm:
            p.paragraph_format.left_indent = Cm(left_indent_cm)
        _render_inline(p, text)

    def equation(self, text: str) -> None:
        text = _sanitize_xml(text or "").strip()
        if not text:
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        # Native display equation (OMML) for the whole expression; if the
        # converter can't handle it, fall back to the inline renderer (which
        # handles $...$ spans + unicode approximation).
        omath = latex_to_omml_element(text)
        if omath is not None:
            p._p.append(omath)
        else:
            _render_inline(p, text)

    def labeled(self, label: str, body: str) -> None:
        """A 'Label: body' paragraph with hanging indent so wrapped lines
        align under body. Used for Question/Options/Answer/Solution.

        If the body contains a Markdown-style table (`| h1 | h2 |` +
        `|---|---|` separator), the table is rendered as a real Word
        table directly below the label paragraph instead of being
        dumped as raw pipe-text.
        """
        chunks = list(_extract_tables_from_text(body.strip()))
        # Label paragraph always exists; first 'text' chunk (if any)
        # becomes the inline body next to the bold label.
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(2.4)
        p.paragraph_format.first_line_indent = Cm(-2.4)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(f"{label}:  ")
        r.bold = True
        r.font.size = Pt(10)
        first_text_consumed = False
        for kind, payload in chunks:
            if kind == "text":
                if not payload.strip():
                    continue
                if not first_text_consumed:
                    _render_inline(p, payload)
                    first_text_consumed = True
                else:
                    cp = self.doc.add_paragraph()
                    cp.paragraph_format.left_indent = Cm(2.4)
                    cp.paragraph_format.space_after = Pt(2)
                    _render_inline(cp, payload)
            else:
                headers, rows = payload
                self.table(headers, rows)

    def options(self, opts: list[tuple[str, str]]) -> None:
        """Render MCQ options. Inline if short, stacked otherwise.

        Each option body goes through `_render_inline` so embedded
        `$...$` math and bare LaTeX commands render the same way
        as Question/Solution bodies — no raw `\\Rightarrow` or `x^2`.
        """
        if not opts:
            return
        inline = all(len(b) <= 28 for _, b in opts)
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(2.4)
        p.paragraph_format.first_line_indent = Cm(-2.4)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("Options:  ")
        r.bold = True
        r.font.size = Pt(10)
        if inline:
            # Build a single string with separators, then route through
            # _render_inline so each option's `$...$` math + LaTeX
            # commands get the same Unicode treatment.
            joined = "    ".join(f"({L}) {b}" for L, b in opts)
            _render_inline(p, joined)
        else:
            # First option on the label line, rest stacked
            for idx, (L, b) in enumerate(opts):
                if idx == 0:
                    _render_inline(p, f"({L}) {b}")
                else:
                    p2 = self.doc.add_paragraph()
                    p2.paragraph_format.left_indent = Cm(2.4)
                    p2.paragraph_format.space_after = Pt(1)
                    _render_inline(p2, f"({L}) {b}")

    def keypoint(self, body: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("Key Point:  ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = KEYPOINT_ORANGE
        _render_inline(p, body)

    def bullets(self, items: list[str]) -> None:
        # Double-pointer fix: list items frequently already carry their own
        # ordinal ("1." / "i." / "(a)"). The "List Bullet" style ALSO adds a
        # marker → "• 1. text". When the source items carry their own
        # numbering, render them as plain indented paragraphs (the source
        # ordinal IS the single marker, numbering preserved); otherwise use
        # the bullet style. Either way: exactly ONE marker per item.
        _ord = re.compile(r"^\s*(?:\(?\d+\)?[.)]|\(?[ivxIVXa-eA-E]\)?[.)])\s+")
        has_ordinals = sum(1 for it in items if _ord.match(str(it))) >= max(
            1, (len(items) + 1) // 2
        )
        for item in items:
            s = str(item).strip()
            if has_ordinals:
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.space_after = Pt(2)
                _render_inline(p, s)
            else:
                p = self.doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                _render_inline(p, s)

    def table(self, headers: list[str], rows: list[list[str]]) -> None:
        if not headers and not rows:
            return
        cols = max(len(headers), max((len(r) for r in rows), default=0))
        if cols == 0:
            return
        tbl = self.doc.add_table(rows=(1 if headers else 0) + len(rows), cols=cols)
        try:
            tbl.style = "Light Grid Accent 1"
        except KeyError:  # built-in style missing on some Word versions
            pass
        row_offset = 0
        if headers:
            for i, h in enumerate(headers):
                cell = tbl.rows[0].cells[i]
                cell.text = ""
                # Route through _render_inline so equation/chem LaTeX in the
                # header renders as a native equation; then bold the text runs.
                _render_inline(cell.paragraphs[0], str(h))
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            row_offset = 1
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                if ci >= cols:
                    continue
                cell = tbl.rows[ri + row_offset].cells[ci]
                cell.text = ""
                _render_inline(cell.paragraphs[0], str(val))

    def question_gap(self) -> None:
        """Visual breather between two questions in the same group."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)

    def section_gap(self) -> None:
        """Bigger breather between adjacent sections."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)

    def figure_callout(self, label: str, caption: str = "") -> None:
        label = _sanitize_xml(label or "")
        caption = _sanitize_xml(caption or "")
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text = f"[Figure: {label}"
        if caption:
            text += f" — {caption}"
        text += "]"
        r = p.add_run(text)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED

    def image(self, image_bytes: bytes, label: str = "", caption: str = "",
              max_width_inches: float = 5.0) -> None:
        """Embed a figure binary in the doc. Centered, capped at
        ``max_width_inches`` so big images don't overflow the page.
        Label + caption render as a centred italic figcaption below."""
        if not image_bytes:
            return
        label = _sanitize_xml(label or "")
        caption = _sanitize_xml(caption or "")
        from docx.shared import Inches
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        try:
            run.add_picture(io.BytesIO(image_bytes), width=Inches(max_width_inches))
        except Exception:
            # Fall back to text callout if python-docx fails to read the bytes
            self.figure_callout(label or "image", caption)
            return
        if label or caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(6)
            if label:
                r1 = cap.add_run(label)
                r1.bold = True
                r1.italic = True
                r1.font.size = Pt(9)
                r1.font.color.rgb = MUTED
            if label and caption:
                r2 = cap.add_run(" — ")
                r2.italic = True
                r2.font.size = Pt(9)
                r2.font.color.rgb = MUTED
            if caption:
                r3 = cap.add_run(caption)
                r3.italic = True
                r3.font.size = Pt(9)
                r3.font.color.rgb = MUTED

    def to_bytes(self) -> bytes:
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()


# ---------------------------------------------------------------------------
# Questions / Regen rendering
# ---------------------------------------------------------------------------

def _render_question_head(
    b: _DocBuilder, q: dict, *, label: str | None = None
) -> None:
    """Render the question STEM + options only.

    Split out so callers can render embedded_figures BETWEEN the stem
    and the solution (matching the PDF's original layout where the
    diagram sits between the problem statement and the worked-out
    steps).
    """
    # NEVER strip anything from raw_text. The full extracted text — stem,
    # options, whatever's in there — goes into the doc verbatim, so no content
    # can ever silently disappear during export. Previously this code stripped
    # options off the stem and only re-rendered them when `q.has_options` was
    # True; when the flag was wrong, options were lost. Removing the strip
    # entirely eliminates that entire class of data-loss bug at the cost of
    # not having a separate "Options:" section header (options appear inline
    # in the stem, which is exactly how they sit in the raw_text). User
    # directive: "no data should be missing".
    stem = (q.get("raw_text") or "").strip()
    question_label = "Question" if not label else f"Question {label}"
    b.labeled(question_label, stem)


def _render_question_tail(b: _DocBuilder, q: dict) -> None:
    """Render Answer / Solution / question gap after the stem (and after
    any embedded_figures the caller emitted between head and tail)."""
    sol_text = q.get("solution_text") or ""
    answer, solution = _split_answer_from_solution(sol_text)
    if answer:
        b.labeled("Answer", answer)
    if solution:
        b.labeled("Solution", solution)
    b.question_gap()


def _maybe_embed_regen_diagram(b: _DocBuilder, q: dict) -> bool:
    """Step 2 — embed the regenerated LaTeX/SVG diagram (rasterized to PNG)
    when the question carries one and it is NOT a fallback.

    Returns True when a diagram image was embedded, so the caller can SKIP the
    original figure (the new diagram REPLACES it). Returns False — meaning keep
    the original figure — when the feature is off, no diagram is present, the
    model fell back to the original, the SVG is empty, or rasterization failed.
    """
    if not settings.EMBED_REGEN_DIAGRAM_IN_DOCX:
        return False
    rd = q.get("regenerated_diagram")
    if not isinstance(rd, dict) or rd.get("fallback_to_original"):
        return False
    svg = (rd.get("svg_preview") or "").strip()
    if not svg:
        return False
    from app.services.svg_raster import rasterize_svg_to_png

    png = rasterize_svg_to_png(svg)
    if not png:
        # Neither cairosvg nor resvg could render it — keep the original figure.
        return False
    subject = (rd.get("subject") or "").strip()
    label = "Regenerated diagram" + (f" · {subject}" if subject else "")
    b.image(png, label=label, caption="")
    return True


def _render_question(b: _DocBuilder, q: dict, *, label: str | None = None) -> None:
    """Backward-compatible single-call render — stem + options + answer
    + solution + gap. Used by the question-bank export path which has no
    embedded_figures handling. The final-draft path uses
    _render_question_head + figures + _render_question_tail so figures
    sit between stem and solution."""
    _render_question_head(b, q, label=label)
    # Step 2 — embed the regenerated diagram between stem and solution when
    # present (no-op for bank questions, which never carry one).
    _maybe_embed_regen_diagram(b, q)
    _render_question_tail(b, q)


def build_questions_docx(
    title: str,
    sections: list[dict[str, Any]],
    *,
    section_only: str | None = None,
) -> bytes:
    """Build the bank-export Word doc.

    `sections` shape: list of {section_ref, section_title, questions: [...]}.
    When `section_only` is set, restrict to that section_ref and skip the
    cover title (single-section export).
    """
    b = _DocBuilder()
    if not section_only:
        b.title(f"{title} — Question Bank")

    # Partition sections into "single example" vs "grouped" buckets so
    # the doc reads naturally: WORKED EXAMPLES heading then each example,
    # then any non-example sections under their own heading.
    example_secs: list[dict] = []
    other_secs: list[dict] = []
    for s in sections:
        if section_only and s.get("section_ref") != section_only:
            continue
        title_l = (s.get("section_title") or "").lower()
        if "example" in title_l:
            example_secs.append(s)
        else:
            other_secs.append(s)

    if example_secs:
        b.group_header("Worked Examples")
        for s in example_secs:
            qs = s.get("questions") or []
            if not qs:
                continue
            b.section_heading(s.get("section_title") or s.get("section_ref") or "")
            for q in qs:
                _render_question(b, q)
            b.section_gap()

    for s in other_secs:
        qs = s.get("questions") or []
        if not qs:
            continue
        b.group_header(s.get("section_title") or s.get("section_ref") or "")
        for idx, q in enumerate(qs, start=1):
            qnum = (q.get("question_number") or "").strip() or str(idx)
            _render_question(b, q, label=qnum)
        b.section_gap()

    return b.to_bytes()


def build_regen_docx(
    label_text: str,
    custom_instructions: str | None,
    sections: list[dict[str, Any]],
    *,
    section_only: str | None = None,
) -> bytes:
    """Build the regen-export Word doc — same layout as questions, plus
    a header note for custom instructions when present."""
    b = _DocBuilder()
    if not section_only:
        b.title(f"{label_text} — Regenerated Questions")
        if custom_instructions:
            p = b.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            r = p.add_run(f"⚡ Custom instructions: {custom_instructions}")
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = MUTED

    # Re-use the questions partitioning logic
    example_secs: list[dict] = []
    other_secs: list[dict] = []
    for s in sections:
        if section_only and s.get("section_ref") != section_only:
            continue
        title_l = (s.get("section_title") or "").lower()
        if "example" in title_l:
            example_secs.append(s)
        else:
            other_secs.append(s)

    if example_secs:
        b.group_header("Worked Examples (Regenerated)")
        for s in example_secs:
            qs = s.get("questions") or []
            if not qs:
                continue
            b.section_heading(s.get("section_title") or s.get("section_ref") or "")
            for q in qs:
                _render_question(b, q)
            b.section_gap()

    for s in other_secs:
        qs = s.get("questions") or []
        if not qs:
            continue
        b.group_header(s.get("section_title") or s.get("section_ref") or "")
        for idx, q in enumerate(qs, start=1):
            qnum = (q.get("question_number") or "").strip() or str(idx)
            _render_question(b, q, label=qnum)
        b.section_gap()

    return b.to_bytes()


# ---------------------------------------------------------------------------
# Theory rendering
# ---------------------------------------------------------------------------

def _render_theory_block(b: _DocBuilder, block: dict, section_title_key: str,
                         last_sub: list[str]) -> None:
    t = block.get("t")
    c = block.get("c", "")
    if t in ("h2", "h3", "h4"):
        if _norm(c) == section_title_key:
            return  # adjacent duplicate of section heading
        if _norm(c) == _norm(last_sub[0]):
            return
        b.sub_heading(c)
        last_sub[0] = c
    elif t == "p":
        b.paragraph(c)
    elif t == "eq":
        b.equation(c)
    elif t == "example":
        # Worked example inside theory body
        lbl = (block.get("label") or "Example").strip()
        prob = (block.get("prob") or "").strip()
        sol = (block.get("sol") or "").strip()
        b.sub_heading(lbl)
        if prob:
            b.paragraph(prob, left_indent_cm=0.4)
        if sol:
            p = b.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run("Solution:  ")
            r.bold = True
            r.font.size = Pt(10)
            _render_inline(p, sol)
    elif t in ("kp", "remember", "keypoint"):
        b.keypoint(c)
    elif t == "list":
        items = block.get("items") or []
        b.bullets(items)
    elif t == "table":
        headers = block.get("headers") or []
        rows = block.get("rows") or []
        if not headers and not rows:
            # Real data stores tables as raw `\begin{tabular}` in `c` (no
            # structured headers/rows) → parse it so it renders as a Word
            # table instead of being silently dropped.
            ph, pr = _parse_latex_tabular(c)
            if ph is not None:
                headers, rows = ph, pr
        b.table(headers, rows)
    elif t == "figure":
        b.figure_callout(block.get("label") or "unlabelled",
                         block.get("caption") or "")
    else:
        if c:
            b.paragraph(c)


def build_theory_docx(
    book_title: str,
    sections: list[dict[str, Any]],
) -> bytes:
    """Build the theory Word doc.

    `sections` shape: list of {section_id, title, blocks: [...]}.
    """
    b = _DocBuilder()
    b.title(book_title)
    for s in sections:
        title = s.get("title") or s.get("section_id") or ""
        blocks = s.get("blocks") or []
        if not blocks:
            continue
        b.section_heading(title)
        title_key = _norm(title)
        last_sub = [""]
        for blk in blocks:
            _render_theory_block(b, blk, title_key, last_sub)
        b.section_gap()
    return b.to_bytes()


# ---------------------------------------------------------------------------
# Final Draft rendering — Phase 3 Composer output
# ---------------------------------------------------------------------------
# Walks the FinalDraft.items ordered list and routes each item to the
# right _DocBuilder method, reusing the polished formatting that
# build_theory_docx + build_questions_docx use:
#   - Section headings styled consistently
#   - Theory blocks (p/h3/eq/def/kp/list/table/example) routed through
#     _render_theory_block so paragraphs, equations, key-points, lists,
#     tables, and worked-examples all look identical to the standalone
#     theory export.
#   - Questions routed through _render_question so MCQs split options,
#     solutions get the right typography, embedded figures embed as
#     images via the new _DocBuilder.image method.
#   - Figures embed binaries (not text callouts) at section level.
#   - Custom text blocks render as plain paragraphs.

def _render_paragraph_with_tables(b: _DocBuilder, text: str) -> None:
    """Render a paragraph that MAY contain embedded pipe-tables. Splits
    on the table markers via `_extract_tables_from_text` so each table
    becomes a real Word table; surrounding prose becomes plain paragraphs.
    Used for `p` blocks (which often carry OCR'd value-tables inline)."""
    if not text or not text.strip():
        return
    chunks = list(_extract_tables_from_text(text.strip()))
    if not any(k == "table" for k, _ in chunks):
        b.paragraph(text.strip())
        return
    for kind, payload in chunks:
        if kind == "text":
            if payload.strip():
                b.paragraph(payload.strip())
        else:  # table
            headers, rows = payload
            b.table(headers, rows)


def _render_block_item(b: _DocBuilder, block: dict[str, Any],
                       title_key: str, last_sub: list[str]) -> None:
    """Like _render_theory_block but tolerates the v2 block types the
    Composer might receive (kp / def / *_ref). Also extracts inline
    pipe-tables from `p` blocks so they render as real Word tables."""
    t = block.get("t")
    if t == "p":
        _render_paragraph_with_tables(b, block.get("c") or "")
        return
    if t == "def":
        # Definition: term in bold, body underneath
        term = (block.get("term") or "").strip()
        body = (block.get("c") or "").strip()
        if term:
            b.sub_heading(f"Definition — {term}")
        if body:
            b.paragraph(body, left_indent_cm=0.4)
        return
    if t in ("example_ref", "exercise_ref", "question_ref"):
        # A3 fix — SUPPRESS unmatched chips in the final DOCX. Matched chips
        # have already been removed from the blocks list by
        # `_merge_chips_with_questions` (their target question gets inlined
        # at the chip's position). Any chip surviving to this point is
        # either orphan (no matching question anywhere) or garbage bled
        # into a non-question section by theory over-extraction. Emitting
        # "→ Exercise: N" placeholders for those clutters the document
        # (e.g. Shortcuts polluted with 74 empty chips). The reader can
        # always inspect the source PDF if a question is genuinely
        # missing. To restore the old "spot me" rendering, look up
        # commit 58c5e1e in git history.
        return
    if t == "fig":
        # The seeder (seed_draft_items_from_merge) drops fig BLOCKS when
        # a matching figure ITEM exists at the same position. A fig block
        # that survives to this renderer means the embedder could NOT
        # link a figure here — emit a muted figure_callout so the
        # exported document shows that a figure was expected at this
        # spot, rather than a silent gap.
        c = (block.get("c") or "").strip()
        if c:
            b.figure_callout("Figure", c)
        return
    # Delegate to the existing theory renderer for the standard types
    _render_theory_block(b, block, title_key, last_sub)


def _render_custom_text(b: _DocBuilder, content: str) -> None:
    """Render a user-added custom_text item. Splits on blank lines so
    each paragraph is its own block. Inline markdown ($math$, bold) is
    handled by _render_inline which mirrors what paragraphs use."""
    if not content:
        return
    for chunk in re.split(r"\n\s*\n", content.strip()):
        chunk = chunk.strip()
        if not chunk:
            continue
        # Route through paragraph() so markdown headings (#, ##, ###) and
        # inline bold/italic/code render as real Word formatting.
        b.paragraph(chunk, space_after_pt=6)


def _question_sort_key(question_number: Any) -> tuple:
    """Sort key for `question_number`. Returns (numeric_prefix, full_string).

    Handles all the formats the extractor produces:
      "1"     -> (1, "1")
      "5"     -> (5, "5")
      "5(a)"  -> (5, "5(a)")
      "5(b)"  -> (5, "5(b)")
      "5(i)"  -> (5, "5(i)")
      "10"    -> (10, "10")
      "Q5"    -> (5, "Q5")
      ""/None -> (10**9, "")   missing — sort to end

    Secondary string comparison groups variants of the same parent together
    in alpha/roman order: 5 < 5(a) < 5(b) < 5(i) < 6.
    """
    s = str(question_number or "").strip()
    if not s:
        return (10**9, "")
    m = re.search(r"\d+", s)
    num = int(m.group()) if m else 10**9
    return (num, s)


def _sort_question_runs(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Return a copy of `items` with consecutive `question` items sorted by
    their `question_number`. Non-question items (section_heading, block,
    figure, custom_text) keep their original positions — only the order
    within a contiguous question run is changed.

    The upstream order (from the composer / API) is by `created_at`, which
    is extractor-time and effectively random within a page. This sort
    presents questions in textbook-original numeric order on every
    rendering surface that calls this helper.

    Render-layer only — the underlying API / DB / composer queries are
    untouched (per user constraint: "i don't want to touch any api it
    might break few things"). Same logic will be mirrored in the frontend
    Preview / Composer pages.
    """
    out: list[dict[str, Any]] = []
    i = 0
    while i < len(items):
        if items[i].get("type") == "question":
            j = i
            while j < len(items) and items[j].get("type") == "question":
                j += 1
            run = list(items[i:j])
            run.sort(key=lambda it: _question_sort_key(
                (it.get("question") or {}).get("question_number")
            ))
            out.extend(run)
            i = j
        else:
            out.append(items[i])
            i += 1
    return out


def build_final_draft_docx(
    book_title: str,
    items: list[dict[str, Any]],
    figure_bytes_map: dict[str, bytes],
) -> bytes:
    """Walk the FinalDraft items list and render a polished Word doc.

    ``figure_bytes_map`` is keyed by figure_id (str). The caller is
    responsible for materialising the image bytes (variant=regen if
    approved, else original) before invoking this function so this
    module stays sync-only.
    """
    # Sort questions within each contiguous run by question_number so the
    # textbook's original numeric order is preserved on export. The upstream
    # items list is in extractor-creation order (effectively random within a
    # page), which causes the "Q1, Q3, Q2, Q6, Q5, Q4..." jumble.
    items = _sort_question_runs(items)

    b = _DocBuilder()
    b.title(book_title or "Final Draft")

    current_section_key = ""
    last_sub = [""]
    # Per-section question counter — resets at every section_heading. Used as
    # the label for "Question N:" when the extractor didn't populate
    # q.question_number (or left it empty). Without this, the final-draft path
    # was the ONLY export path that didn't pass `label` to _render_question_head
    # — every question rendered as just "Question:" with no number. The bank
    # export paths (build_questions_docx / build_regen_docx) already had this
    # pattern; we just bring the final-draft path in line.
    section_q_counter = 0
    for it in items:
        t = it.get("type")
        if t == "section_heading":
            title = (it.get("title") or it.get("section_id") or "").strip()
            # Use section_heading for ALL section items so the typography
            # matches the theory/questions exports (13pt NAVY, normal-case,
            # consistent spacing). The composer's level field is positional
            # info; the visual style stays uniform — Word handles outline
            # hierarchy via paragraph styles, not size escalation.
            b.section_heading(title)
            current_section_key = _norm(title)
            last_sub = [""]
            section_q_counter = 0  # reset per-section counter
            continue
        if t == "block":
            _render_block_item(b, it.get("block") or {}, current_section_key, last_sub)
            continue
        if t == "figure":
            f = it.get("figure") or {}
            fid = str(f.get("figure_id") or "")
            data = figure_bytes_map.get(fid)
            if data:
                b.image(data, label=f.get("label") or "", caption=f.get("caption") or "")
            else:
                b.figure_callout(f.get("label") or "image", f.get("caption") or "")
            continue
        if t == "question":
            q = dict(it.get("question") or {})
            # Per-section question number. Use the extracted `question_number`
            # field when present (textbook-original "Q5", "Q19" etc.). When the
            # extractor didn't populate it (empty/None), fall back to the
            # per-section running index — so labels are always sequential 1, 2,
            # 3 per section regardless of upstream data gaps. Mirrors the bank
            # export's pattern (build_questions_docx / build_regen_docx).
            section_q_counter += 1
            qnum = (str(q.get("question_number") or "")).strip() or str(section_q_counter)
            # Layout: stem (+ options) → figures → solution.
            # Matches the PDF's original layout where the construction
            # diagram sits between the problem statement and the worked-
            # out steps. Previously the figures were emitted AFTER the
            # full question card (stem + options + solution), pushing
            # them past the solution text — wrong position relative to
            # the source PDF.
            _render_question_head(b, q, label=qnum)
            # A regenerated diagram REPLACES the original figures (→ empty list
            # so nothing else emits). Otherwise split the embedded figures by
            # body_target: question-stem figures render UNDER THE STEM (before
            # the solution), solution figures render AFTER THE SOLUTION — so each
            # figure lands in the body it belongs to. NULL body_target defaults
            # to the question side. Matches the review view + Preview/Composer.
            _efs = [] if _maybe_embed_regen_diagram(b, q) else (
                q.get("embedded_figures") or []
            )
            _q_figs = [f for f in _efs if (f.get("body_target") or "question") != "solution"]
            _s_figs = [f for f in _efs if f.get("body_target") == "solution"]

            def _emit_qfigs(_figs: list) -> None:
                for f in _figs:
                    fid = str(f.get("figure_id") or "")
                    data = figure_bytes_map.get(fid)
                    if data:
                        b.image(data, label=f.get("label") or "", caption=f.get("caption") or "")
                    else:
                        b.figure_callout(f.get("label") or "image", f.get("caption") or "")

            _emit_qfigs(_q_figs)           # question-body figures under the stem
            _render_question_tail(b, q)    # solution text
            _emit_qfigs(_s_figs)           # solution-body figures after the solution
            continue
        if t == "custom_text":
            _render_custom_text(b, it.get("content") or "")
            continue
        # Unknown item type → ignore silently
    return b.to_bytes()
